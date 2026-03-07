"""
Document Generator: produces legal DOCX documents.

Uses the Strategy pattern (OCP) so new document types can be added without
modifying DocGenerator. Each strategy:
1. Retrieves reference documents of the same type from the knowledge base (RAG).
2. Prompts Qwen2.5-VL to draft the full document text.
3. Assembles the DOCX with python-docx using formal legal formatting.
"""

from __future__ import annotations

import logging
import tempfile
from abc import ABC, abstractmethod
from pathlib import Path
from uuid import uuid4

import ollama
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from src.agent.rag_chain import RAGChain
from src.config import Settings
from src.domain.errors import DocumentGenerationError
from src.domain.models import ConversationContext, DocumentContext, GeneratedDocument

logger = logging.getLogger(__name__)

_DOC_SYSTEM_PROMPT = """Eres un abogado colombiano experto redactando documentos legales formales.
Redacta el documento completo basándote en:
1. La información del usuario proporcionada.
2. Los documentos de referencia del mismo tipo extraídos de la base de conocimiento.
Usa lenguaje formal y jurídico colombiano. Incluye todos los encabezados, considerandos,
hechos, pretensiones y peticiones apropiados para el tipo de documento.
No uses placeholders — usa "XXXXXX" solo donde sea estrictamente necesario (ej. número de proceso).
"""


class DocumentStrategy(ABC):
    """Base class for document generation strategies."""

    doc_type: str

    @abstractmethod
    async def generate(self, context: DocumentContext, llm: ollama.AsyncClient, model: str) -> str:
        """Generate document text. Returns raw text content."""
        ...

    def get_generation_prompt(self, context: DocumentContext) -> str:
        """Build the user prompt for the LLM."""
        ref_text = "\n\n---\n\n".join(
            f"Referencia [{i + 1}] ({r.doc_type or 'documento'}):\n{r.excerpt}"
            for i, r in enumerate(context.reference_chunks)
        )
        memory_text = (
            "\n".join(f"- {e.key}: {e.value}" for e in context.memory)
            if context.memory
            else "No disponible"
        )

        return (
            f"Tipo de documento a redactar: {self.doc_type.upper()}\n\n"
            f"Información del usuario/caso:\n{memory_text}\n\n"
            f"Solicitud del usuario: {context.user_request}\n\n"
            f"Documentos de referencia del mismo tipo:\n{ref_text}\n\n"
            f"Redacta el {self.doc_type} completo:"
        )


class TutelaStrategy(DocumentStrategy):
    doc_type = "tutela"

    async def generate(self, context: DocumentContext, llm: ollama.AsyncClient, model: str) -> str:
        prompt = self.get_generation_prompt(context)
        response = await llm.chat(
            model=model,
            messages=[
                {"role": "system", "content": _DOC_SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            options={"temperature": 0.3},
        )
        return response.message.content.strip()


class DemandaStrategy(DocumentStrategy):
    doc_type = "demanda"

    async def generate(self, context: DocumentContext, llm: ollama.AsyncClient, model: str) -> str:
        prompt = self.get_generation_prompt(context)
        response = await llm.chat(
            model=model,
            messages=[
                {"role": "system", "content": _DOC_SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            options={"temperature": 0.3},
        )
        return response.message.content.strip()


class DerechoPeticionStrategy(DocumentStrategy):
    doc_type = "derecho_de_peticion"

    async def generate(self, context: DocumentContext, llm: ollama.AsyncClient, model: str) -> str:
        prompt = self.get_generation_prompt(context)
        response = await llm.chat(
            model=model,
            messages=[
                {"role": "system", "content": _DOC_SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            options={"temperature": 0.3},
        )
        return response.message.content.strip()


class DocxAssembler:
    """Converts raw text into a formatted DOCX with legal styling."""

    @staticmethod
    def assemble(text: str, doc_type: str) -> bytes:
        """Build DOCX bytes from raw text."""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

        title = doc.add_heading(doc_type.upper().replace("_", " "), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(14)
        title_run.bold = True

        doc.add_paragraph()

        paragraphs = text.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.isupper() or (len(para_text) < 80 and para_text.endswith(":")):
                p = doc.add_heading(para_text, level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p = doc.add_paragraph(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        doc.save(str(tmp_path))
        content = tmp_path.read_bytes()
        tmp_path.unlink(missing_ok=True)
        return content


class DocGenerator:
    """
    Orchestrates document generation using registered strategies.

    OCP: add new strategies without modifying this class.
    """

    def __init__(
        self,
        strategies: list[DocumentStrategy],
        rag_chain: RAGChain,
        llm_client: ollama.AsyncClient,
        settings: Settings,
        output_dir: Path,
    ) -> None:
        self._registry: dict[str, DocumentStrategy] = {s.doc_type: s for s in strategies}
        self._rag = rag_chain
        self._llm = llm_client
        self._settings = settings
        self._output_dir = output_dir
        self._output_dir.mkdir(parents=True, exist_ok=True)
        self._assembler = DocxAssembler()

    def supported_types(self) -> list[str]:
        return list(self._registry.keys())

    async def generate(
        self,
        doc_type: str,
        user_request: str,
        ctx: ConversationContext,
    ) -> GeneratedDocument:
        """
        Generate a DOCX document:
        1. Find strategy for doc_type (fall back to generic if unknown).
        2. RAG-retrieve references of the same type.
        3. Strategy generates text via LLM.
        4. Assemble DOCX.
        5. Persist to disk and return metadata.
        """
        normalized = doc_type.lower().replace(" ", "_")
        strategy = self._registry.get(normalized)
        if strategy is None:
            strategy = self._registry.get("tutela") or next(iter(self._registry.values()))
            logger.warning(
                "No strategy for '%s', using '%s' as fallback", doc_type, strategy.doc_type
            )

        rag_result = await self._rag.query(
            question=f"estructura y formato de {doc_type}",
            ctx=ctx,
            doc_type_filter=normalized,
        )
        reference_chunks = rag_result.sources

        doc_ctx = DocumentContext(
            doc_type=normalized,
            conversation_id=ctx.conversation_id,
            user_request=user_request,
            reference_chunks=reference_chunks,
            memory=ctx.memory,
        )

        try:
            text = await strategy.generate(doc_ctx, self._llm, self._settings.chat_model)
        except Exception as exc:
            raise DocumentGenerationError(
                f"LLM failed to generate {doc_type}: {exc}", doc_type=doc_type
            ) from exc

        try:
            docx_bytes = DocxAssembler.assemble(text, doc_type)
        except Exception as exc:
            raise DocumentGenerationError(
                f"DOCX assembly failed: {exc}", doc_type=doc_type
            ) from exc

        doc_id = uuid4()
        filename = f"{normalized}_{doc_id.hex[:8]}.docx"
        output_path = self._output_dir / filename
        output_path.write_bytes(docx_bytes)

        return GeneratedDocument(
            document_id=doc_id,
            doc_type=normalized,
            filename=filename,
            download_url=f"/api/v1/documents/{doc_id}/download",
        )
